from hashlib import new
from operator import truediv
from tkinter import *
from tkinter import filedialog
import tkinter
from docx import Document
import re
import copy
from docx import document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import docx
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import os
from Class_write_excel import *
from Class_Convert_katakana import *
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from datetime import datetime
import shutil
from Class_Read_Code import *
import docx2txt
import win32com
from win32com import client
self_database = read_file_json(".\\Database\\Database.json")
_define_meaning, _define_name, _define_value, _define_file = self_database["Define_meaning"], self_database["Define_name"],self_database["Define_value"],self_database["Define_file"]


##### gom data của 1 dic thành data chuẩn
def extract_data(dic):
    for data in dic:
        for _data in data.items():
            yield _data[0],_data[1]

class Word_:
    def __init__(self,links = "",linkscode = "",new = False,accept_trackchange=False) -> None:
        self.path_DD = links
        self.name_file = linkscode.replace("/","\\").split("\\")[-1]
        self.document_teamplet = Document("Document_teamplet.docx")
        self.Table_teamplet_enum     = self.document_teamplet.tables[0]
        self.Table_teamplet_struct   = self.document_teamplet.tables[1]
        self.Table_teamplet_g_vari   = self.document_teamplet.tables[2]
        self.Table_teamplet_c_vari   = self.document_teamplet.tables[3]
        self.Table_teamplet_define   = self.document_teamplet.tables[4]
        self.Paragraphs_teamplet = self.document_teamplet.paragraphs[0]
        self.new = new
        self.accept_trackchange = accept_trackchange
        if self.new == False and accept_trackchange == False:
            self.docx = Session_DD(self.path_DD)
            self.docx.read_Data_table_definition()
        if accept_trackchange == True:
            self.document = Document(self.path_DD)
        else:
            self.document = Document(self.docx.path_file_AcceptAll_change) if self.new == False else Document()
        self.DD_name = self.name_file.split(".")[0]
        # self.Excel_Data = Read_excel("Database.xlsx")
        self.define_meaning, self.define_name, self.define_value, self.define_file = _define_meaning, _define_name, _define_value, _define_file

    @staticmethod
    def convert_paragraph_full_size(paragraph):
        for texxt in paragraph._element.xpath('.//w:t'):
            texxt.text = Convert(texxt.text).convert_haft2full_katakana()
        
    @staticmethod
    def Creat_table_common(document, table, para_Heading = True, paragraph = False, text_for_paragraph = "None", level_heading = 3):
        paragraph = document.add_paragraph()
        if para_Heading == True:
            paragraph = document.add_heading(level=level_heading)
        elif paragraph == True:
            paragraph = document.add_paragraph()
        run = paragraph.add_run(text_for_paragraph)
        font = run.font
        font.color.rgb = RGBColor(0,0,0)

        tbl, p = table._tbl, paragraph._p
        new_tbl = copy.deepcopy(tbl)
        p.addnext(new_tbl)
        return document.tables[len(document.tables)-1]

    @staticmethod
    def insert_row(table,count_row):
        current_row = table.rows[count_row] 
        table.rows[count_row].height_rule = WD_ROW_HEIGHT_RULE.AUTO
        tbl = table._tbl
        border_copied = copy.deepcopy(current_row._tr)
        tr = border_copied
        current_row._tr.addnext(tr)
        return 1


    def creat_enum(self,Dic_enum):
        self.document.add_heading(level=1).add_run("Enum")
        for data_key, data_value in extract_data(Dic_enum):
            row_inset = 5
            Word_.Creat_table_common(self.document,self.Table_teamplet_enum,text_for_paragraph = data_key[0])
            table = self.document.tables[len(self.document.tables)-1]
            table.cell(0,1).text = data_key[1]
            table.cell(1,1).text = data_key[2].replace("enum","").strip()
            for data in range(len(data_value)):
                if data > 2:
                    self.insert_row(table,row_inset)
                    row_inset += 1
                table.cell(3+data,1).text = str(data_value[data][0])
                table.cell(3+data,2).text = str(data_value[data][1])
                table.cell(3+data,3).text = str(data_value[data][2])

    def creat_Struct(self,Dic_struct):
        self.document.add_heading(level=1).add_run("Struct")
        for data_key, data_value in extract_data(Dic_struct):
            row_inset = 5
            Word_.Creat_table_common(self.document,self.Table_teamplet_struct,text_for_paragraph = data_key[0])
            table = self.document.tables[len(self.document.tables)-1]
            table.cell(0,1).text = data_key[1]
            table.cell(1,1).text = data_key[2].replace("struct","").strip()
            for data in range(len(data_value)):
                if data > 2:
                    self.insert_row(table,row_inset)
                    row_inset += 1
                table.cell(3+data,1).text = str(data_value[data][0])
                table.cell(3+data,2).text = str(data_value[data][1])
                table.cell(3+data,3).text = str(data_value[data][2])

    def  creat_define(self,list_define):
        self.document.add_heading(level=1).add_run("Define")
        for data_value in list_define:
            Word_.Creat_table_common(self.document,self.Table_teamplet_define,text_for_paragraph = data_value[0])
            table = self.document.tables[len(self.document.tables)-1]
            table.cell(0,1).text = data_value[1]
            table.cell(2,1).text = str(data_value[2])

    def creat_g_variable(self,list_variable):
        self.document.add_heading(level=1).add_run("Global Variable")
        Variable_File = "ファイルローカル"
        Variable_Global = "グローバル"
        for data_value in list_variable:
            Word_.Creat_table_common(self.document,self.Table_teamplet_g_vari,text_for_paragraph = data_value[2])
            table = self.document.tables[len(self.document.tables)-1]
            table.cell(0,1).text = data_value[0] + re.sub("[\w\*]+","",data_value[1]).replace(" ","")
            table.cell(1,1).text = data_value[1].replace("[","").replace("]","")
            table.cell(2,1).text = data_value[0] + "_Range_QA"
            table.cell(4,1).text = self.read_initial_value(data_value[4])
            table.cell(5,1).text = Variable_File if data_value[4].find("static") != -1 else Variable_Global
            # table.cell(6,1).text = self.extract_remark_array(data_value[3])
            data_remark = self.extract_remark_array(data_value[3])
            if len(data_remark) == 1:
                table.cell(6,1).text = data_remark[0]
            else:
                table.cell(6,1).text = data_remark[0]
                for anotion in data_remark[1]:
                    table.cell(6,1).add_paragraph(anotion)


    @staticmethod
    def read_initial_value(string):
        string = re.sub("\/\*[\s\S]*?\*\/|\/\/.*","",string)
        match = re.search("=(.+?);",string)
        if match:
            return match.group(1).strip()
        else:
            return "-"

    @staticmethod
    def extract_len_array(string):
        links_source_code = r"C:\Projects\PMAS\PMA-15-RE\PMA-15-RE"
        remark = "配列サイズ = [DMA配列最大値(※1)]"
        string_cscope = "cscope -R -L -1 "
        match = re.search("\[(.*?)\]",string)
        if match:
            data_array = re.findall("\[(.*?)\]",string)
            for data in data_array:
                os.chdir(links_source_code)
                if Word_.cmd_and_read(string_cscope + data) == "cscope: cannot find file -1":
                    print("xxxx")
        else:
            return "-"

    def extract_remark_array(self, string):

        list_not_anotion = ["define.h"]
        list_interface = ["if_mdl.c","if_mdl.h","if_rte.c","if_rte.h","if_mdl_fs.c","if_mdl_fs.h","if_mdl_sys.c","if_mdl_sys.h","if_knl.c","if_knl.h","if_ibl_knl.c","if_ibl_knl.h","if_mdl_dcdc.h","if_mdl_dcdc.c"]
        index_anotion = {}

        list_add_or_no = []        

        remark_exemple = "配列サイズ = [DMA配列最大値(※1)]"
        remark = "配列サイズ = ["
        string_anotion = ""
        match = re.search("\[(.*?)\]",string)
        if match:
            data_array_tam = re.findall("\[(.*?)\]",string)
            data_array = []
            for data_tam in data_array_tam:
                for dt in re.findall("[\w]+",data_tam):
                    data_array.append(dt)

            list_remark = []
            list_anotion = []
            for i in range(len(data_array)):
                string = data_array[i].strip()
                if string == "":
                    continue
                string_tam,string_anotion_tam, index_anotion = self.add_remark(string,index_anotion,list_not_anotion,list_interface)
                list_remark.append(string_tam)
                string = re.sub("\(\s*※\s*\d+\s*\)","",string_anotion_tam)
                if string not in list_add_or_no:
                    list_anotion.append(string_anotion_tam)
                list_add_or_no.append(string)
            string_remark_tam = " * ".join(list_remark) + "]"
            return [remark + string_remark_tam, list_anotion]
        else:
            return ["-"]

    def add_remark(self, string, index_anotion, list_not_anotion, list_interface):

        string_interface = "インタフェース設計書[file_name]を参照"
        string_file = "SW詳細設計書[file_name]を参照"
        string_in_file = "マクロ定義を参照"
        anotion = "※"
        string_anotion = ""

        index = None
        if self.define_name.count(string) > 1:
            for i in range(len(self.define_name)):
                if self.define_name[i] == string and self.define_file[i] == self.name_file:
                    index = i
                    break

        if index != None:
            file = self.define_file[index]
            mening = str(self.define_meaning[index])
        else:
            try:
                index = self.define_name.index(string)
            except:
                return string, string_anotion, index_anotion
            file = self.define_file[index]
            mening = str(self.define_meaning[index])

        remark = ""
        if file not in list_not_anotion:
            ipos = ""
            if file in list_interface:
                file = file.split(".")[0].replace("if_","").upper()
                string_interface = string_interface.replace("file_name",file)
                ipos = string_interface
                if string_interface not in index_anotion:
                    index_anotion[string_interface] = len(index_anotion) + 1
                string_anotion = "(" + anotion + str(index_anotion[ipos]) + ")" + string_interface
            elif file == self.name_file:
                if string_in_file not in index_anotion:
                    index_anotion[string_in_file] = len(index_anotion) + 1
                ipos = string_in_file
                string_anotion = "(" + anotion + str(index_anotion[ipos]) + ")" + string_in_file
            else:
                file = file.split(".")[0]
                string_file = string_file.replace("file_name",file)
                if string_file not in index_anotion:
                    index_anotion[string_file] = len(index_anotion) + 1
                ipos = string_file
                string_anotion = "(" + anotion + str(index_anotion[ipos]) + ")" + string_file
            
            remark = mening + "(" + anotion + str(index_anotion[ipos]) + ")"
        else:
            remark = mening
        return remark, string_anotion, index_anotion

    @staticmethod
    def cmd_and_read(command):
        return os.popen(command).read()

    """Đọc đoạn test bất kỳ có thuộc tính paragraphs bỏ qua được trackchange"""
    @staticmethod
    def convert_paragraphs_2_text(paragraph,all_text = True):
        rs = paragraph._element.xpath('.//w:t')
        if all_text != True:  # dùng cho 1 số 1 trường hợp đặc biệt
            string = ""
            for i in range(len(rs)-1):
                if i > 0:
                    string += rs[i].text
            return string
        else:
            return u"".join([r.text for r in rs])
    
    "Đọc hết 1 cell"
    @staticmethod
    def convert_cell_paragraphs_2_text(cell_paragraph):
        string = ""
        for _paragraph in cell_paragraph:
            string += Word_.convert_paragraphs_2_text(_paragraph)
        return string

    @staticmethod
    def change_font_size(paragraph,size_):
        # stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
        # paragraph.clear()
        # paragraph.add_run(stringparagraph)
        # Word_.convert_paragraph_full_size(paragraph)
        stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
        font_name_in_table = "ＭＳ Ｐゴシック"
        font_name_heading = "ＭＳ ゴシック"
        font_name = "Times New Roman"

        for run in paragraph.all_runs:
            run.font.size = docx.shared.Pt(size_)
            run.font.color.rgb = RGBColor(0,0,0)
            if size_ == 9:
                run.font.name = "ＭＳ Ｐゴシック"
            else:
                run.font.name = "ＭＳ ゴシック"
            run.font.name = font_name
            if stringparagraph == "Control number" or stringparagraph == "Version number" or stringparagraph == "System name" or stringparagraph == "Approved By" or stringparagraph == "Verified By" or stringparagraph == "Created By":
                run.bold = TRUE
        return 1

    """bỏ các cell trùng nhau"""
    @staticmethod
    def iter_unique_cells(row):
        prior_tc = None
        for cell in row.cells:
            this_tc = cell._tc
            if this_tc is prior_tc:
                continue
            prior_tc = this_tc
            yield cell

    """inter gom tuần tự paragraph->table theo thứ tự trong word"""
    @staticmethod
    def iter_block_items(object):
        object_ = object
        if isinstance(object_, _Document):
            parent_elm = object_.element.body
        elif isinstance(object_, _Cell):
            parent_elm = object_._tc
        elif isinstance(object_, _Row): # chưa chạy được
            parent_elm = object_._tr
        else:
            raise ValueError("Có gì đó không đúng")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, object_)
            elif isinstance(child, CT_Tbl):
                yield Table(child, object_)

    def Change_font_paragraph_all(self):
        change_font = False
        for paragraph in self.document.paragraphs:
            if change_font == True:
                stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
                if stringparagraph == "":
                    continue
                outputuFF = re.findall(r'[\uFF61-\uFF9F]', stringparagraph)
                if len(outputuFF) > 0:
                    # paragraph.text = Convert(stringparagraph).convert_haft2full_katakana()
                    self.convert_paragraph_full_size(paragraph)
                if paragraph.header_level == 1:
                    Word_.change_font_size(paragraph,14)
                elif paragraph.header_level == 2:
                    Word_.change_font_size(paragraph,12)
                else:
                    Word_.change_font_size(paragraph,10.5)
            elif paragraph.header_level == 1:
                stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
                outputuFF = re.findall(r'[\uFF61-\uFF9F]', stringparagraph)
                if len(outputuFF) > 0:
                    # paragraph.text = Convert(stringparagraph).convert_haft2full_katakana()
                    self.convert_paragraph_full_size(paragraph)
                change_font = True
                Word_.change_font_size(paragraph,14)

    def Change_font_paragraph_table_all(self):
        cnt_table = 0
        for table in self.document.tables:
            cnt_table += 1
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        for _table in cell.tables:
                            for _row in _table.rows:
                                for _cell in Word_.iter_unique_cells(_row):
                                    if _cell.tables:  
                                        for __table in _cell.tables:
                                            for __row in __table.rows:
                                                for __cell in Word_.iter_unique_cells(__row):
                                                    for paragraph in __cell.paragraphs:
                                                        Word_._check_katakana(paragraph)
                                                        Word_.change_font_size(paragraph,9)
                                    for paragraph in _cell.paragraphs:
                                        Word_._check_katakana(paragraph)
                                        Word_.change_font_size(paragraph,9)
                    for paragraph in cell.paragraphs:
                        stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
                        if stringparagraph == "":
                            continue
                        outputuFF = re.findall(r'[\uFF61-\uFF9F]', stringparagraph)
                        if len(outputuFF) > 0:
                            # paragraph.text = Convert(stringparagraph).convert_haft2full_katakana()
                            self.convert_paragraph_full_size(paragraph)
                        if cnt_table != 1:
                            Word_.change_font_size(paragraph,9)
                        else:
                            Word_.change_font_size(paragraph,9)

    def fix_allWord_(self):
        self.Change_font_paragraph_all()
        self.Change_font_paragraph_table_all()

    def _KatakanaConverter(self):
        for paragraph in self.document.paragraphs:
            Word_._check_katakana(paragraph)
        cnt_table = 0
        for table in self.document.tables:
            cnt_table += 1
            if cnt_table > 2:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.tables:
                            for _table in cell.tables:
                                for _row in _table.rows:
                                    for _cell in Word_.iter_unique_cells(_row):
                                        if _cell.tables:  
                                            for __table in _cell.tables:
                                                for __row in __table.rows:
                                                    for __cell in Word_.iter_unique_cells(__row):
                                                        for paragraph in __cell.paragraphs:
                                                            Word_._check_katakana(paragraph)
                                        for paragraph in _cell.paragraphs:
                                            Word_._check_katakana(paragraph)
                        for paragraph in cell.paragraphs:
                            Word_._check_katakana(paragraph)

    @staticmethod
    def _check_katakana(paragraph):
        stringparagraph = Word_.convert_paragraphs_2_text(paragraph)
        outputuFF = re.findall(r'[\uFF61-\uFF9F]', stringparagraph)
        if len(outputuFF) > 0:
            # paragraph.text = Convert(stringparagraph).convert_haft2full_katakana()
            Word_.convert_paragraph_full_size(paragraph)
            if isinstance(paragraph._parent,docx.table._Cell):
                Word_.change_font_size(paragraph,9)
            elif paragraph.header_level == 1:
                Word_.change_font_size(paragraph,14)
            elif paragraph.header_level == 2:
                Word_.change_font_size(paragraph,12)
            else:
                Word_.change_font_size(paragraph,10.5)
            # if Word_.add_comment == True:
            #     paragraph.add_comment("Updated  haftwith to full width", author="Nhan Bui")

    def save(self):
        print("______________________Saving DD_______________________")
        self.fix_allWord_()
        if self.new == True:
            if self.name_file.find(".h") != -1:
                self.DD_name = self.DD_name + "_h"
            self.document.save(self.DD_name + ".docx")
            path_save = os.getcwd() + "\\" + self.DD_name + ".docx"
            self.save_txt(path_save)
            strCurrentPath_update_toc = os.getcwd() + "\\temp"

            if not os.path.exists(strCurrentPath_update_toc):
                string_mkdir = 'mkdir ' + '"'+ strCurrentPath_update_toc + '"'
                os.popen(string_mkdir).read()
            shutil.move(path_save, strCurrentPath_update_toc + "\\" + self.DD_name + ".docx")
            shutil.move(path_save.replace(".docx",".txt"), strCurrentPath_update_toc + "\\" + self.DD_name + ".txt")
            print(strCurrentPath_update_toc + "\\" + self.DD_name + ".docx")

        else:
            if self.accept_trackchange:
                self.document.save(self.path_DD)
                print(self.path_DD)
            else:
                self.document.save(self.docx.path_file_AcceptAll_change)
                print(self.docx.path_file_AcceptAll_change)
        print("______________________Save DD Complete_______________________")

    def save_txt(self,path_file_docx):
        strCurrentPath_update_toc = path_file_docx
        text = docx2txt.process(strCurrentPath_update_toc)
        lines = text.split("\n")
        non_empty_lines = [line for line in lines if line.strip() != ""]
        string_without_empty_lines = "\n".join(non_empty_lines)
        text = string_without_empty_lines
        strCurrentPath_update_toc = strCurrentPath_update_toc.replace("docx","txt")
        with open(strCurrentPath_update_toc, "w",encoding="UTF-8") as text_file:
            print(text, file=text_file)
            text_file.close()
        print(strCurrentPath_update_toc)

    @staticmethod
    def copy_table_in_DD(table, block, text_Paragraph = "", level = 3, Paragraphs_teample = None):
        paragraph_none = Document().add_paragraph()
        paragraph_none.text = ""
        paragraph_none_tam = paragraph_none._p
        new_p_none = copy.deepcopy(paragraph_none_tam)

        paragraph_none2 = Document().add_paragraph()
        paragraph_none2.text = ""
        paragraph_none_tam2 = paragraph_none2._p
        new_p_none2 = copy.deepcopy(paragraph_none_tam2)

        paragraph = copy.deepcopy(Paragraphs_teample)
        paragraph.style = Document().styles['Heading '+str(level)]
        try:
            paragraph._element.xpath('.//w:t')[0].text = text_Paragraph
        except:
            paragraph.text = text_Paragraph

        for run in paragraph.all_runs:
            font = run.font
            font.color.rgb = RGBColor(0,0,0)

        p = paragraph._p
        new_p = copy.deepcopy(p)

        if isinstance(block, Paragraph):
            tbl, block_2 = table._tbl, block._p
            new_tbl = copy.deepcopy(tbl)
            block_2.addnext(new_p)
            new_p.addnext(new_tbl)
            # new_tbl.addnext(new_p_none2)

        elif isinstance(block, Table):
            tbl, block_2 = table._tbl, block._tbl
            new_tbl = copy.deepcopy(tbl)
            block_2.addnext(new_p_none)
            new_p_none.addnext(new_p)
            new_p.addnext(new_tbl)
            # new_tbl.addnext(new_p_none2)

        return new_tbl
    
    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def trackchange_Paragraphs_cell(cell, Paragraphs_teample,text_Paragraph = ""):
        paragraph = copy.deepcopy(Paragraphs_teample)
        try:
            paragraph._element.xpath('.//w:t')[0].text = text_Paragraph
        except:
            paragraph.text = text_Paragraph

        for run in paragraph.all_runs:
            font = run.font
            font.color.rgb = RGBColor(0,0,0)

        cell.text = ""
        cell.paragraphs[0]._p.addnext(paragraph._p)
        Word_.delete_paragraph(cell.paragraphs[0])
        return cell

    def insert_DD_enum(self,Dic_enum):
        element_index = 0
        name_compare = ""
        for data_key, data_value in extract_data(Dic_enum):
            if element_index ==0  and data_key[1] not in self.docx.data["Enum"]:
                for paragraph in self.document.paragraphs:
                    if  paragraph.header_level == 2 and Word_.convert_paragraphs_2_text(paragraph).find("Enumeration definition") != -1:
                        table_tam = self.copy_table_in_DD(self.Table_teamplet_enum,paragraph,data_key[0],Paragraphs_teample = self.Paragraphs_teamplet)
                        self.check_table(table_tam,data_key,data_value)
                    else:
                        if Word_.convert_paragraphs_2_text(paragraph) == "本書において定義される列挙体はない。":
                            paragraph.clear()
            elif data_key[1] not in self.docx.data["Enum"]:
                for itable in range(len(self.document.tables)):
                    try:
                        if Word_.convert_paragraphs_2_text(self.document.tables[itable].cell(0,1)) == name_compare:
                            table_tam = self.copy_table_in_DD(self.Table_teamplet_enum,self.document.tables[itable],data_key[0],Paragraphs_teample = self.Paragraphs_teamplet)
                            self.check_table(table_tam,data_key,data_value,self.document.tables[itable+1])
                            break
                    except:
                        pass
            name_compare = data_key[1]
            element_index += 1
    def check_table(self,table_teamplet,data_key,data_value,Table = None,enum_struct = None):
        if Table != None:
            self.modify_table_enum(Table,data_key,data_value,enum_struct)
            return None
        for table in self.document.tables:
            if table_teamplet == table._tbl:
                self.modify_table_enum(table,data_key,data_value,enum_struct)
                return None
    def modify_table_enum(self,table,data_key,data_value,enum_struct = None):
        row_inset = 5
        self.trackchange_Paragraphs_cell(table.cell(0,1),self.Paragraphs_teamplet,data_key[1])
        string_enum_struct = data_key[2].replace("enum","").strip() if enum_struct == None else data_key[2].replace("struct","").strip()
        self.trackchange_Paragraphs_cell(table.cell(1,1),self.Paragraphs_teamplet,string_enum_struct) 
        for data in range(len(data_value)):
            if data > 2:
                self.insert_row(table,row_inset)
                row_inset += 1
            self.trackchange_Paragraphs_cell(table.cell(3+data,1),self.Paragraphs_teamplet,str(data_value[data][0]))
            self.trackchange_Paragraphs_cell(table.cell(3+data,2),self.Paragraphs_teamplet,str(data_value[data][1]))
            self.trackchange_Paragraphs_cell(table.cell(3+data,3),self.Paragraphs_teamplet,str(data_value[data][2]))

    def insert_DD_struct(self,Dic_struct):
        element_index = 0
        name_compare = ""
        for data_key, data_value in extract_data(Dic_struct):
            if element_index ==0  and data_key[1] not in self.docx.data["Struct"]:
                for paragraph in self.document.paragraphs:
                    if  paragraph.header_level == 2 and Word_.convert_paragraphs_2_text(paragraph).find("Structure definition") != -1:
                        table_tam = self.copy_table_in_DD(self.Table_teamplet_struct,paragraph,data_key[0],Paragraphs_teample = self.Paragraphs_teamplet)
                        self.check_table(table_tam,data_key,data_value,enum_struct=1)
                    else:
                        if Word_.convert_paragraphs_2_text(paragraph) == "本書において定義される構造体はない。":
                            paragraph.clear()
            elif data_key[1] not in self.docx.data["Struct"]:
                for itable in range(len(self.document.tables)):
                    try:
                        if Word_.convert_paragraphs_2_text(self.document.tables[itable].cell(0,1)) == name_compare:
                            table_tam = self.copy_table_in_DD(self.Table_teamplet_struct,self.document.tables[itable],data_key[0],Paragraphs_teample = self.Paragraphs_teamplet)
                            self.check_table(table_tam,data_key,data_value,self.document.tables[itable+1],1)
                            break
                    except:
                        pass
            name_compare = data_key[1]
            element_index += 1

    def insert_DD_define(self,list_define):
        element_index = 0
        name_compare = ""
        for data_value in list_define:
            if element_index == 0  and data_value[1] not in self.docx.data["Macro"]:
                for paragraph in self.document.paragraphs:
                    if  paragraph.header_level == 2 and Word_.convert_paragraphs_2_text(paragraph).find("Macro definition") != -1:
                        table_tam = self.copy_table_in_DD(self.Table_teamplet_define,paragraph,data_value[0],Paragraphs_teample = self.Paragraphs_teamplet)
                        self.check_table_4x_3_4_5(table_tam,data_value,x4_3_4_5=5)
                    else:
                        if Word_.convert_paragraphs_2_text(paragraph) == "本書において定義されるマクロはない。":
                            paragraph.clear()
            elif data_value[1] not in self.docx.data["Macro"]:
                for itable in range(len(self.document.tables)):
                    try:
                        if Word_.convert_paragraphs_2_text(self.document.tables[itable].cell(0,1)) == name_compare:
                            table_tam = self.copy_table_in_DD(self.Table_teamplet_define,self.document.tables[itable],data_value[0],Paragraphs_teample = self.Paragraphs_teamplet)
                            self.check_table_4x_3_4_5(table_tam,data_value,self.document.tables[itable+1],5)
                            break
                    except:
                        pass
            
            name_compare = data_value[1]
            element_index += 1

    def check_table_4x_3_4_5(self,table_teamplet,data_value,Table = None,x4_3_4_5 = None):
        if Table != None:
            self.modify_table_x4_3_4_5(Table,data_value,x4_3_4_5)
            return None
        for table in self.document.tables:
            if table_teamplet == table._tbl:
                self.modify_table_x4_3_4_5(table,data_value,x4_3_4_5)
                return None

    def modify_table_x4_3_4_5(self,table,data_value,x4_3_4_5):
        if x4_3_4_5 == 5:  #cho bảng define
            self.trackchange_Paragraphs_cell(table.cell(0,1),self.Paragraphs_teamplet,str(data_value[1]))
            self.trackchange_Paragraphs_cell(table.cell(2,1),self.Paragraphs_teamplet,str(data_value[2]))
        elif x4_3_4_5 == 3:  #cho bảng biếng global và static in file
            Variable_File = "ファイルローカル"
            Variable_Global = "グローバル"
            string_name_variable = data_value[0] + re.sub("[\w\*]+","",data_value[1]).replace(" ","")
            string_tydef_variable = data_value[1].replace("[","").replace("]","")
            string_range_variable = data_value[0] + "_Range_QA"
            string_init_variable = self.read_initial_value(data_value[4])
            string_scpoe_variable = Variable_File if data_value[4].find("static") != -1 else Variable_Global
            data_remark = self.extract_remark_array(data_value[3])

            self.trackchange_Paragraphs_cell(table.cell(0,1),self.Paragraphs_teamplet,string_name_variable)
            self.trackchange_Paragraphs_cell(table.cell(1,1),self.Paragraphs_teamplet,string_tydef_variable)
            self.trackchange_Paragraphs_cell(table.cell(2,1),self.Paragraphs_teamplet,string_range_variable)
            self.trackchange_Paragraphs_cell(table.cell(4,1),self.Paragraphs_teamplet,string_init_variable)
            self.trackchange_Paragraphs_cell(table.cell(5,1),self.Paragraphs_teamplet,string_scpoe_variable)
            self.trackchange_Paragraphs_cell(table.cell(6,1),self.Paragraphs_teamplet,data_remark[0])
            if len(data_remark) > 1:
                ipos = 0
                for anotion in data_remark[1]:
                    try:
                        self.trackchange_Paragraphs_addnext(table.cell(6,1),self.Paragraphs_teamplet,anotion,ipos)
                        ipos += 1
                    except:
                        table.cell(6,1).add_paragraph(anotion)
            
    def insert_DD_variable(self,list_variable):
        element_index = 0
        name_compare = ""
        for data_value in list_variable:
            string_name_variable = data_value[0] + re.sub("[\w\*]+","",data_value[1]).replace(" ","")
            if element_index == 0  and string_name_variable not in self.docx.data["Global"]:
                for paragraph in self.document.paragraphs:
                    if  paragraph.header_level == 2 and Word_.convert_paragraphs_2_text(paragraph).find("Global variable definition") != -1:
                        table_tam = self.copy_table_in_DD(self.Table_teamplet_g_vari,paragraph,data_value[2],Paragraphs_teample = self.Paragraphs_teamplet)
                        self.check_table_4x_3_4_5(table_tam,data_value,x4_3_4_5=3)
                    else:
                        if Word_.convert_paragraphs_2_text(paragraph) == "本書において定義されるグローバル変数はない。":
                            paragraph.clear()
            elif string_name_variable not in self.docx.data["Global"]:
                for itable in range(len(self.document.tables)):
                    try:
                        if Word_.convert_paragraphs_2_text(self.document.tables[itable].cell(0,1)) == name_compare:
                            table_tam = self.copy_table_in_DD(self.Table_teamplet_g_vari,self.document.tables[itable],data_value[2],Paragraphs_teample = self.Paragraphs_teamplet)
                            self.check_table_4x_3_4_5(table_tam,data_value,self.document.tables[itable+1],3)
                            break
                    except:
                        pass
            name_compare = string_name_variable
            element_index += 1

    @staticmethod
    def trackchange_Paragraphs_addnext(cell, Paragraphs_teample,text_Paragraph = "",ipos=0):
        paragraph = copy.deepcopy(Paragraphs_teample)
        try:
            paragraph._element.xpath('.//w:t')[0].text = text_Paragraph
        except:
            paragraph.text = text_Paragraph
        for run in paragraph.all_runs:
            font = run.font
            font.color.rgb = RGBColor(0,0,0)
        try:
            cell.paragraphs[ipos]._p.addnext(paragraph._p)
        except:
            cell.paragraphs[-1]._p.addnext(paragraph._p)
        return cell


class Session_DD:
    def __init__(self,links_Word) -> None:
        links_Word = links_Word.replace("/","\\")
        self.path_file = links_Word
        self.data = {
                        "Enum": {},
                        "Struct": {},
                        "Global": {},
                        "Constant": {},
                        "Macro": {},
                    }
        self.path_file_AcceptAll_change = ""

    @staticmethod
    def cmd_and_read(command):
        return os.popen(command).read()

    def AcceptAll_change(self,path):
        file_path = path.replace("/","\\")
        file_path = self.cmd_and_read("Accept_all_change.exe " + '"' + file_path + '"')
        # try:
        #     word = client.gencache.EnsureDispatch('Word.Application')
        # except AttributeError:
        #     # Remove cache and try again.
        #     MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        #     for module in MODULE_LIST:
        #         if re.match(r'win32com\.gen_py\..+', module):
        #             del sys.modules[module]
        #     shutil.rmtree(os.path.abspath(os.path.join(win32com.__gen_path__, '..')))
        #     word = client.gencache.EnsureDispatch('Word.Application')


        # word.Visible = False
        # doc = word.Documents.Open(file_path )
        # # Accept all revisions
        # word.ActiveDocument.Revisions.AcceptAll()
        # word.ActiveDocument.Save()
        # doc.Close(False)
        # word.Application.Quit()

        strCurrentPath_update_toc = os.getcwd() + "\\temp\\DD_accept_trackchange"

        if not os.path.exists(strCurrentPath_update_toc):
            string_mkdir = 'mkdir ' + '"'+ strCurrentPath_update_toc + '"'
            os.popen(string_mkdir).read()

        name_file = path.replace("/","\\").split("\\")[-1]
        path_accept_trackchange = strCurrentPath_update_toc
        path_accept_trackchange += "\\" + name_file  
        if os.path.exists(path_accept_trackchange):
            os.remove(path_accept_trackchange)
        shutil.move(file_path, path_accept_trackchange)

        self.path_file_AcceptAll_change = path_accept_trackchange
        return path_accept_trackchange

    @staticmethod
    def convert_paragraphs_2_text(paragraph):
        rs = paragraph._element.xpath('.//w:t')
        return u"".join([r.text for r in rs])

    """inter gom tuần tự paragraph->table theo thứ tự trong word"""
    @staticmethod
    def iter_block_items(object):
        object_ = object
        if isinstance(object_, _Document):
            parent_elm = object_.element.body
        elif isinstance(object_, _Cell):
            parent_elm = object_._tc
        else:
            raise ValueError("Có gì đó không đúng")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, object_)
            elif isinstance(child, CT_Tbl):
                yield Table(child, object_)

    """bỏ các cell trùng nhau"""
    @staticmethod
    def iter_unique_cells(row):
        prior_tc = None
        for cell in row.cells:
            this_tc = cell._tc
            if this_tc is prior_tc:
                continue
            prior_tc = this_tc
            yield cell

    @staticmethod
    def read_table_column_3_copy(table):
        if len(table.rows[0].cells) == 1:
            return ""
        for row in table.rows:
            for cell in Session_DD.iter_unique_cells(row):
                string = Session_DD.convert_paragraphs_2_text(cell)
                string = re.sub("\(\s*※\s*\d+\s*\)","",string)
                if string == "列挙体名称" or string == "構造体名称" or string == "名称":
                    continue
                return string


    def read_Data_table_definition(self):
        document = Document(self.AcceptAll_change(self.path_file))
        string = ""
        string_meaning = ""
        string_pos = ""
        read = False
        list_data = [ [] for i in range(5)]

        for block in self.iter_block_items(document):
            if isinstance(block, Paragraph):
                string = self.convert_paragraphs_2_text(block)
                if string == "":
                    continue
                else: 
                    string_meaning = string
                if block.style.name == 'Heading 2':
                    if string.find("Enumeration definition") != -1 :
                        read = True
                        string_pos = "Enumeration"
                    elif string.find("Structure definition") != -1:
                        string_pos = "Structure"
                    elif string.find("Global variable definition") != -1:
                        string_pos = "Global"
                    elif string.find("Constant definition") != -1:
                        string_pos = "Constant"
                    elif string.find("Macro definition") != -1:
                        string_pos = "Macro"
                elif string.find("Function definition") != -1 and block.style.name == 'Heading 1':
                    read = False

            elif isinstance(block, Table):
                if read == True:
                    string_cell = self.read_table_column_3_copy(block)
                    if string_pos == "Enumeration":
                        list_data[0].append(string_cell)
                    elif string_pos == "Structure":
                        list_data[1].append(string_cell)
                    elif string_pos == "Global":
                        list_data[2].append(string_cell)
                    elif string_pos == "Constant":
                        list_data[3].append(string_cell)
                    elif string_pos == "Macro":
                        list_data[4].append(string_cell)
        self.data["Enum"] = list_data[0]
        self.data["Struct"] = list_data[1]
        self.data["Global"] = list_data[2]
        self.data["Constant"] = list_data[3]
        self.data["Macro"] = list_data[4]

        return self.data

if __name__ == '__main__':
    # document = Word_()
    # dic = [{('* \\brief define Return type for NM functions. Derived from Std_ReturnType.', 'Nm_ReturnType', 'enum'):[['NM_E_OK', 0, 'Function call has been successfully accomplished and returned.'], ['NM_E_NOT_OK', 1, 'Meaning_variable_enum-None'], ['returned because of an internal execution error. */', 2, 'Meaning_variable_enum-None'], ['NM_E_NOT_EXECUTED', 2, 'Function call is not executed.']]},
    # {('Meaning_enum-None', 'Nm_ModeType', 'enum'): [['NM_MODE_BUS_SLEEP', 0, 'Bus-Sleep Mode'], ['NM_MODE_PREPARE_BUS_SLEEP', 1, 'Prepare-Bus Sleep Mode'], ['NM_MODE_SYNCHRONIZE', 2, 'Synchronize Mode'], ['NM_MODE_NETWORK', 
    # 3, 'Network Mode']]
    # }]
    # document.creat_enum(dic)
    # document.creat_Struct(dic)
    # document.save()
    #####################  end thêm hàng #########################
    # root = tkinter.Tk()
    # root.withdraw()
    # file_path = filedialog.askopenfilename()
    # docx = Session_DD(file_path)
    # docx.read_Data_table_definition()
    # print(docx.data["Enum"])
    # print(docx.data["Struct"])
    # print(docx.data["Global"])
    # print(docx.data["Constant"])
    # print(docx.data["Macro"])
    # links = r"C:\Projects\Schaeffler_Phase_3\01_Working\01_INPUT\01_Source_Code\V1\TQ250_BL21.00.00(SVA1T21TH_21B08A)\forCUSTOMER\NidecSWC_Lib\nidec_lib\PF\MDL\DcDcCtrl\ISR\DcDcCtrl_ISR.c"
    # links = r"C:\Projects\Schaeffler_Phase_3\01_Working\01_INPUT\01_Source_Code\V1\TQ250_BL21.00.00(SVA1T21TH_21B08A)\forCUSTOMER\NidecSWC_Lib\nidec_lib\PF\MDL\DcDcCtrl\DcDcPwmDriver\DcDcCtrl_ontime.c"
    links = r"C:\Projects\Schaeffler_Phase_3\01_Working\01_INPUT\01_Source_Code\V1\TQ250_BL21.00.00(SVA1T21TH_21B08A)\forCUSTOMER\NidecSWC_Lib\nidec_lib\PF\MDL\CDD\CurrentControl\fi_pwm_invadj_pe1.c"
    Data = Typerdef(links)
    Data.read_define()
    Data.read_variable()
    Data.read_typef_enum()
    Data.read_typedef_struct()

    # root = tkinter.Tk()
    # root.withdraw()
    # file_path = filedialog.askopenfilename()
    # document = Word_(file_path,links,accept_trackchange = True)
    # # document.insert_DD_enum(Data.Enum_4_1)
    # # document.insert_DD_struct(Data.Struct_4_2)
    # # document.insert_DD_variable(Data.Variable_4_3)
    # # document.insert_DD_define(Data.Macro_4_5)
    # document.save()


    # document = Word_("",links,new=True)
    # document.creat_enum(Data.Enum_4_1)
    # document.creat_Struct(Data.Struct_4_2)
    # document.creat_define(Data.Macro_4_5)
    # document.creat_g_variable(Data.Variable_4_3)
    # document.save()